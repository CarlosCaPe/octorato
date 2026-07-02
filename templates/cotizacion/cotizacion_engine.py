#!/usr/bin/env python3
"""
Generic, brand-neutral quote / proposal (cotizacion) generator.

This is the INVARIANT engine, the "cimientos" in the fachada/cimientos pattern.
It knows how to lay out a professional service quote as a .docx (and .pdf via
LibreOffice) but it carries ZERO branding and ZERO content of its own. Everything
that is brand-specific (company name, colors, font, logo, signature image) or
content-specific (cover text, letter, sections, rates, tables, dates) is supplied
at run time through two config dicts: `brand` and `content`.

Wire your real brand and clausulado from a PRIVATE config (company/ dir or the
client arm), never from this public template. See README.md.

Usage as a library:
    from cotizacion_engine import CotizacionEngine
    engine = CotizacionEngine(brand, content)
    docx_path, pdf_path = engine.build("out/quote.docx")

Usage as a CLI:
    python3 cotizacion_engine.py --config config.example.yaml
    python3 cotizacion_engine.py --config config.example.yaml --out demo/quote.docx --no-pdf

Requires:
    pip install python-docx lxml pyyaml
    sudo apt install -y libreoffice-writer     (only needed for PDF output)
"""

import argparse
import os
import sys
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ══════════════════════════════════════════════════════════════════════════════
# Config helpers
# ══════════════════════════════════════════════════════════════════════════════

def hex_to_rgb(value):
    """Accept '#RRGGBB' or 'RRGGBB' and return a python-docx RGBColor."""
    if isinstance(value, RGBColor):
        return value
    v = str(value).lstrip("#")
    return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))


ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class CotizacionEngine:
    """Renders a brand-neutral service quote from a brand config and a content config."""

    def __init__(self, brand: dict, content: dict, base_dir: Path = None):
        self.brand = brand or {}
        self.content = content or {}
        # base_dir is where relative asset paths (signature image, template) resolve from.
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()

        # Typography, all parameterized (no literal font buried in logic).
        self.font_name = self.brand.get("font", "Calibri")
        sizes = self.brand.get("sizes", {})
        self.body_size = Pt(sizes.get("body", 11))
        self.h2_size = Pt(sizes.get("h2", 15))
        self.h3_size = Pt(sizes.get("h3", 13))
        self.small_size = Pt(sizes.get("small", 9))
        self.cover_title_size = Pt(sizes.get("cover_title", 18))
        self.cover_subtitle_size = Pt(sizes.get("cover_subtitle", 14))
        self.cover_client_size = Pt(sizes.get("cover_client", 22))

        # Palette, all from config (no hardcoded 0x00B0F0 style constants).
        colors = self.brand.get("colors", {})
        self.c_brand = hex_to_rgb(colors.get("brand", "3B82F6"))
        self.c_dark = hex_to_rgb(colors.get("dark", "1E3A8A"))
        self.c_black = hex_to_rgb(colors.get("black", "000000"))
        self.c_gray = hex_to_rgb(colors.get("gray", "595659"))
        self.c_white = hex_to_rgb(colors.get("white", "FFFFFF"))
        self.c_pending = hex_to_rgb(colors.get("pending", "C00000"))
        # Table header fill + alternating row fill, hex strings for OOXML shading.
        self.header_fill = str(colors.get("table_header", "1F497D")).lstrip("#")
        self.row_fill = str(colors.get("table_row_alt", "F2F2F2")).lstrip("#")
        self.about_label_fill = str(colors.get("about_label", "E8EDF3")).lstrip("#")

    # ── Low-level builders (invariant format machinery) ───────────────────────

    def clear_body(self, doc):
        body = doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
        return doc

    def add_paragraph(self, doc, text="", style=None, bold=False, italic=False,
                      font_size=None, color=None, alignment=None,
                      space_before=None, space_after=None, font_name=None,
                      keep_with_next=False):
        font_name = font_name or self.font_name
        p = doc.add_paragraph(style=style)
        if alignment is not None:
            p.alignment = alignment
        fmt = p.paragraph_format
        if space_before is not None:
            fmt.space_before = space_before
        if space_after is not None:
            fmt.space_after = space_after
        if keep_with_next:
            fmt.keep_with_next = True
        if text:
            run = p.add_run(text)
            run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = color
        return p

    def add_run(self, paragraph, text, bold=False, italic=False, font_size=None,
                color=None, font_name=None):
        font_name = font_name or self.font_name
        run = paragraph.add_run(text)
        run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
        return run

    def add_page_break(self, doc):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        return p

    def set_table_borders(self, table, color="000000", sz=4):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.append(tblPr)
        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(parse_xml(borders_xml))

    def set_cell_shading(self, cell, color_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_text(self, cell, text, bold=False, font_size=None, color=None,
                      alignment=None, font_name=None):
        font_name = font_name or self.font_name
        font_size = font_size or self.body_size
        color = color if color is not None else self.c_black
        cell.text = ""
        p = cell.paragraphs[0]
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = font_size
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
        return p

    def _set_cell_text_with_inline_bold(self, cell, text, font_size=None, color=None,
                                        font_name=None):
        """Render **bold** spans inside a table cell."""
        import re
        font_name = font_name or self.font_name
        font_size = font_size or self.body_size
        color = color if color is not None else self.c_black
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = font_name
            run.font.size = font_size
            if color:
                run.font.color.rgb = color
            if i % 2 == 1:
                run.bold = True
        return p

    def add_table(self, doc, headers, rows, col_widths=None, header_color=None):
        header_color = header_color or self.header_fill
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(table)

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self.set_cell_shading(cell, header_color)
            self.set_cell_text(cell, header, bold=True, font_size=Pt(10),
                               color=self.c_white, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for r, row_data in enumerate(rows):
            for c, cell_text_val in enumerate(row_data):
                cell = table.rows[r + 1].cells[c]
                if "**" in cell_text_val:
                    self._set_cell_text_with_inline_bold(cell, cell_text_val, font_size=Pt(10))
                else:
                    self.set_cell_text(cell, cell_text_val, bold=False, font_size=Pt(10))
                if r % 2 == 1:
                    self.set_cell_shading(cell, self.row_fill)

        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    row.cells[i].width = width
        return table

    def _col_widths(self, spec):
        """Turn a list of cm numbers from config into Cm() widths, or None."""
        if not spec:
            return None
        return [Cm(w) for w in spec]

    def update_footer_dates(self, doc, new_date):
        """Rewrite the Author-bound footer date SDT and nudge footer anchors.

        This is invariant format machinery. It only matters when a base template
        .docx with a themed footer is supplied. On a blank Document it is a no-op.
        """
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        w_ns = ns["w"]
        wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

        doc.core_properties.author = new_date

        for sect in doc.sections:
            for ftr_type, footer in [("regular", sect.footer),
                                     ("first_page", sect.first_page_footer)]:
                ftr_elem = footer._element

                for sdt in ftr_elem.findall(".//w:sdt", ns):
                    alias = sdt.find(".//w:sdtPr/w:alias", ns)
                    if alias is not None and alias.get(f"{{{w_ns}}}val") == "Author":
                        parent = sdt.getparent()
                        while parent is not None and parent.tag != f"{{{w_ns}}}p":
                            parent = parent.getparent()
                        if parent is not None:
                            sdt_rPr = sdt.find(f".//{{{w_ns}}}rPr")
                            parent.remove(sdt)
                            for child in list(parent):
                                if child.tag != f"{{{w_ns}}}pPr":
                                    parent.remove(child)
                            r_date = etree.SubElement(parent, f"{{{w_ns}}}r")
                            rPr = etree.SubElement(r_date, f"{{{w_ns}}}rPr")
                            if sdt_rPr is not None:
                                for ch in sdt_rPr:
                                    rPr.append(ch.__deepcopy__(True))
                            else:
                                sz = etree.SubElement(rPr, f"{{{w_ns}}}sz")
                                sz.set(f"{{{w_ns}}}val", "24")
                            color_elem = rPr.find(f"{{{w_ns}}}color")
                            if color_elem is None:
                                color_elem = etree.SubElement(rPr, f"{{{w_ns}}}color")
                            color_elem.set(f"{{{w_ns}}}val", "000000")
                            t_date = etree.SubElement(r_date, f"{{{w_ns}}}t")
                            t_date.text = new_date

                for anchor in ftr_elem.findall(f".//{{{wp_ns}}}anchor"):
                    posV = anchor.find(f"{{{wp_ns}}}positionV")
                    if posV is None:
                        continue
                    rel = posV.get("relativeFrom", "")
                    extent = anchor.find(f"{{{wp_ns}}}extent")
                    if extent is None:
                        continue
                    cx = int(extent.get("cx", "0"))
                    cy = int(extent.get("cy", "0"))

                    if cx > 4572000 and cy < 91440 and rel == "bottomMargin":
                        align_elem = posV.find(f"{{{wp_ns}}}align")
                        if align_elem is not None and align_elem.text == "top":
                            posV.remove(align_elem)
                            offset = etree.SubElement(posV, f"{{{wp_ns}}}posOffset")
                            offset.text = "228600"

                    elif cx < 2000000 and cy > 200000:
                        bodyPr = anchor.find(f".//{{{a_ns}}}bodyPr")
                        if bodyPr is not None:
                            bodyPr.set("tIns", "0")
                        align_elem = posV.find(f"{{{wp_ns}}}align")
                        if align_elem is not None:
                            posV.remove(align_elem)
                        offset_elem = posV.find(f"{{{wp_ns}}}posOffset")
                        if offset_elem is None:
                            offset_elem = etree.SubElement(posV, f"{{{wp_ns}}}posOffset")
                        pgnum_offset = "18288" if ftr_type == "first_page" else "-36576"
                        offset_elem.text = pgnum_offset

    # ── Shared sections (invariant house format, filled from content) ─────────

    def _signature_image_path(self):
        sig = self.brand.get("signature_image")
        if not sig:
            return None
        p = Path(sig)
        if not p.is_absolute():
            p = self.base_dir / p
        return p if p.exists() else None

    def build_cover_page(self, doc):
        cov = self.content.get("cover", {})
        for _ in range(3):
            self.add_paragraph(doc, "", space_before=Pt(0), space_after=Pt(0))

        for line in cov.get("title_lines", []):
            self.add_paragraph(doc, line, font_size=self.cover_title_size,
                               color=self.c_dark, space_after=Pt(0))
        subtitle = cov.get("subtitle")
        if subtitle:
            self.add_paragraph(doc, subtitle, font_size=self.cover_subtitle_size,
                               color=self.c_brand, space_after=Pt(36))
        else:
            self.add_paragraph(doc, "", space_after=Pt(36))

        for _ in range(4):
            self.add_paragraph(doc, "", space_before=Pt(0), space_after=Pt(0))

        p = self.add_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        self.add_run(p, cov.get("client_name", ""), bold=True,
                     font_size=self.cover_client_size, color=self.c_dark)

        for _ in range(4):
            self.add_paragraph(doc, "", space_before=Pt(0), space_after=Pt(0))

        prepared_label = cov.get("prepared_for_label", "Prepared for")
        self.add_paragraph(doc, prepared_label, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           font_size=self.body_size, color=self.c_gray, space_after=Pt(0))
        for line in cov.get("prepared_for", []):
            self.add_paragraph(doc, line, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                               font_size=self.body_size, space_after=Pt(0))

        created_label = cov.get("created_by_label", "Created by")
        self.add_paragraph(doc, created_label, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           font_size=self.body_size, color=self.c_gray,
                           space_before=Pt(12), space_after=Pt(0))
        for line in cov.get("created_by", []):
            self.add_paragraph(doc, line, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                               font_size=self.body_size, space_after=Pt(0))

        for _ in range(3):
            self.add_paragraph(doc, "", space_before=Pt(0), space_after=Pt(0))

        date_long = self.content.get("document", {}).get("proposal_date_long", "")
        if date_long:
            self.add_paragraph(doc, date_long, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                               font_size=self.body_size, color=self.c_gray, space_after=Pt(0))

    def build_intro_letter(self, doc):
        letter = self.content.get("intro_letter")
        if not letter:
            return
        p = self.add_paragraph(doc, "", space_after=Pt(0))
        self.add_run(p, letter.get("addressee_bold", ""), bold=True, font_size=self.body_size)
        for line in letter.get("addressee_lines", []):
            self.add_paragraph(doc, line, font_size=self.body_size, space_after=Pt(0))
        self.add_paragraph(doc, "")
        salutation = letter.get("salutation")
        if salutation:
            self.add_paragraph(doc, salutation, font_size=self.body_size, space_after=Pt(12))

        # The signature-line paragraph may highlight the company name in brand color.
        company = self.brand.get("company", "")
        for para in letter.get("paragraphs", []):
            if company and isinstance(para, str) and "{{company}}" in para:
                pre, post = para.split("{{company}}", 1)
                pp = self.add_paragraph(doc, "", font_size=self.body_size,
                                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                if pre:
                    self.add_run(pp, pre, font_size=self.body_size)
                self.add_run(pp, company, font_size=self.body_size, color=self.c_brand, bold=True)
                if post:
                    self.add_run(pp, post, font_size=self.body_size)
            else:
                self.add_paragraph(doc, para, font_size=self.body_size,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))

        closing = letter.get("closing", "")
        if closing:
            self.add_paragraph(doc, closing, font_size=self.body_size, space_after=Pt(24))

        self._insert_signature_image(doc)
        self.add_paragraph(doc, self.brand.get("founder", ""), font_size=self.body_size,
                           bold=True, space_before=Pt(0), space_after=Pt(0))
        self.add_paragraph(doc, self.brand.get("founder_title", ""), font_size=self.body_size,
                           space_before=Pt(0), space_after=Pt(0))

    def _insert_signature_image(self, doc):
        sig = self._signature_image_path()
        if sig:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(str(sig), width=Inches(1.2))
        else:
            self.add_paragraph(doc, "[signature]", font_size=self.small_size,
                               color=self.c_gray, space_after=Pt(2))

    def build_about_page(self, doc):
        about = self.content.get("about")
        if not about:
            return
        heading = about.get("heading", f"About {self.brand.get('company', '')}").strip()
        self.add_paragraph(doc, heading, font_size=self.h2_size, bold=True,
                           color=self.c_dark, space_before=Pt(18), space_after=Pt(8),
                           keep_with_next=True)

        company = self.brand.get("company", "")
        intro = about.get("intro", "")
        if intro:
            if company and "{{company}}" in intro:
                pre, post = intro.split("{{company}}", 1)
                pp = self.add_paragraph(doc, "", font_size=self.body_size,
                                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                if pre:
                    self.add_run(pp, pre, font_size=self.body_size)
                self.add_run(pp, company, font_size=self.body_size, color=self.c_brand, bold=True)
                if post:
                    self.add_run(pp, post, font_size=self.body_size)
            else:
                self.add_paragraph(doc, intro, font_size=self.body_size,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))

        rows = about.get("rows", [])
        if rows:
            table = doc.add_table(rows=len(rows), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self.set_table_borders(table)
            widths = about.get("col_widths_cm", [4, 13])
            for r, (label, value) in enumerate(rows):
                self.set_cell_text(table.rows[r].cells[0], label, bold=True,
                                   font_size=Pt(10), color=self.c_dark)
                self.set_cell_shading(table.rows[r].cells[0], self.about_label_fill)
                self.set_cell_text(table.rows[r].cells[1], value, font_size=Pt(10))
                table.rows[r].cells[0].width = Cm(widths[0])
                table.rows[r].cells[1].width = Cm(widths[1])

    def build_validity(self, doc):
        validity = self.content.get("validity")
        if not validity:
            return
        heading = validity.get("heading", "Validity")
        self.add_paragraph(doc, heading, font_size=self.h2_size, bold=True,
                           color=self.c_dark, space_before=Pt(18), space_after=Pt(8),
                           keep_with_next=True)
        self.add_paragraph(doc, validity.get("text", ""), font_size=self.body_size,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12))

    def build_signature_block(self, doc):
        sig = self.content.get("signature")
        if not sig:
            return
        self.add_paragraph(doc, "", space_after=Pt(12))
        self._insert_signature_image(doc)

        p = self.add_paragraph(doc, "", space_before=Pt(0), space_after=Pt(0))
        self.add_run(p, "____________________________", font_size=self.body_size, color=self.c_gray)
        self.add_paragraph(doc, sig.get("name", self.brand.get("founder", "")),
                           bold=True, font_size=self.body_size, space_before=Pt(2), space_after=Pt(0))
        org_line = sig.get("org_line", self.brand.get("founder_title", ""))
        self.add_paragraph(doc, org_line, font_size=self.body_size,
                           space_before=Pt(0), space_after=Pt(0))
        proposal_date = self.content.get("document", {}).get("proposal_date", "")
        if proposal_date:
            self.add_paragraph(doc, proposal_date, font_size=self.body_size,
                               color=self.c_gray, space_before=Pt(0), space_after=Pt(24))

        # Footer line: company in brand color + the rest in gray.
        p = self.add_paragraph(doc, "", font_size=self.small_size, color=self.c_gray,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_run(p, self.brand.get("company", ""), font_size=self.small_size,
                     color=self.c_brand, bold=True)
        tail = sig.get("footer_line", "")
        if tail:
            self.add_run(p, tail, font_size=self.small_size, color=self.c_gray)

    def build_footer(self, doc):
        """Footer legal / copyright note, filled from content['footer']."""
        footer = self.content.get("footer")
        if not footer:
            return
        self.add_paragraph(doc, "", space_before=Pt(12), space_after=Pt(0))
        self.add_paragraph(doc, footer.get("text", ""), font_size=Pt(8),
                           color=self.c_gray, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_after=Pt(0))

    # ── Generic content sections (fully data-driven) ──────────────────────────

    def render_block(self, doc, block):
        """Render one content block. Type dispatches the layout."""
        btype = block.get("type", "paragraph")

        if btype == "paragraph":
            self.add_paragraph(doc, block.get("text", ""), font_size=self.body_size,
                               bold=block.get("bold", False), italic=block.get("italic", False),
                               alignment=ALIGN.get(block.get("align", "justify")),
                               space_after=Pt(block.get("space_after", 6)))

        elif btype == "lead":
            # A paragraph that starts with a colored bold label, then body text.
            p = self.add_paragraph(doc, "", font_size=self.body_size,
                                   alignment=ALIGN.get(block.get("align", "justify")),
                                   space_after=Pt(block.get("space_after", 6)))
            self.add_run(p, block.get("label", ""), bold=True,
                         font_size=self.body_size, color=self.c_dark)
            self.add_run(p, block.get("body", ""), font_size=self.body_size)

        elif btype == "subheading":
            self.add_paragraph(doc, block.get("text", ""), font_size=self.h3_size,
                               bold=True, color=self.c_dark, space_before=Pt(8),
                               space_after=Pt(4), keep_with_next=True)

        elif btype == "bullet":
            p = self.add_paragraph(doc, "", font_size=self.body_size,
                                   space_before=Pt(1), space_after=Pt(4),
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            p.paragraph_format.left_indent = Cm(block.get("indent_cm", 0.6))
            self.add_run(p, "• ", font_size=self.body_size)
            label = block.get("label", "")
            if label:
                self.add_run(p, label, bold=True, font_size=self.body_size, color=self.c_dark)
            self.add_run(p, block.get("body", ""), font_size=self.body_size)

        elif btype == "table":
            self.add_table(doc, block.get("headers", []), block.get("rows", []),
                           col_widths=self._col_widths(block.get("col_widths_cm")))
            self.add_paragraph(doc, "", space_after=Pt(4))

        elif btype == "spacer":
            self.add_paragraph(doc, "", space_after=Pt(block.get("space_after", 6)))

        elif btype == "note":
            self.add_paragraph(doc, block.get("text", ""), font_size=Pt(10),
                               italic=block.get("italic", True),
                               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               space_after=Pt(block.get("space_after", 6)))
        else:
            raise ValueError(f"Unknown block type: {btype}")

    def build_section(self, doc, section):
        heading = section.get("heading")
        if heading:
            self.add_paragraph(doc, heading, font_size=self.h2_size, bold=True,
                               color=self.c_dark, space_before=Pt(18), space_after=Pt(8),
                               keep_with_next=True)
        for block in section.get("blocks", []):
            self.render_block(doc, block)

    # ── Assembly ──────────────────────────────────────────────────────────────

    def _open_document(self):
        template = self.brand.get("template_docx")
        if template:
            tp = Path(template)
            if not tp.is_absolute():
                tp = self.base_dir / tp
            if tp.exists():
                doc = Document(str(tp))
                self.clear_body(doc)
                footer_date = self.content.get("document", {}).get("footer_date")
                if footer_date:
                    self.update_footer_dates(doc, footer_date)
                return doc
            print(f"  note: template_docx not found at {tp}, using a blank document")
        doc = Document()
        try:
            style = doc.styles["Normal"]
            style.font.name = self.font_name
            style.font.size = self.body_size
        except Exception:
            pass
        return doc

    def build(self, out_path, make_pdf=True):
        """Assemble the whole document and (optionally) render a PDF."""
        doc = self._open_document()

        self.build_cover_page(doc)
        self.add_page_break(doc)

        self.build_intro_letter(doc)
        self.build_about_page(doc)
        self.add_page_break(doc)

        sections = self.content.get("sections", [])
        for i, section in enumerate(sections):
            if section.get("page_break_before") and i > 0:
                self.add_page_break(doc)
            self.build_section(doc, section)

        # Closing house sections.
        self.build_validity(doc)
        self.build_signature_block(doc)
        self.build_footer(doc)

        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        print(f"  DOCX: {out_path} ({out_path.stat().st_size:,} bytes)")

        pdf_path = None
        if make_pdf:
            pdf_path = convert_to_pdf(out_path, out_path.parent)
        return out_path, pdf_path


# ══════════════════════════════════════════════════════════════════════════════
# PDF conversion (LibreOffice headless, best effort)
# ══════════════════════════════════════════════════════════════════════════════

def convert_to_pdf(docx_path: Path, out_dir: Path):
    docx_path = Path(docx_path)
    out_dir = Path(out_dir)
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    cmds = [
        ["libreoffice", "--headless",
         f"-env:UserInstallation=file:///tmp/lo_profile_{os.getpid()}",
         "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        ["soffice",
         f"-env:UserInstallation=file:///tmp/lo_profile_{os.getpid()}b",
         "--headless", "--convert-to", "pdf",
         "--outdir", str(out_dir), str(docx_path)],
        ["flatpak-spawn", "--host", "soffice",
         "--headless", "--convert-to", "pdf",
         "--outdir", str(out_dir), str(docx_path)],
    ]
    for cmd in cmds:
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            if result.returncode == 0 and pdf_path.exists():
                print(f"  PDF:  {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    print(f"  note: LibreOffice not available, no PDF generated for {docx_path.name}")
    return None


# ══════════════════════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════════════════════

def load_config(config_path: Path):
    import yaml
    with open(config_path, "r", encoding="utf-8") as fh:
        cfg = yaml.safe_load(fh)
    brand = cfg.get("brand", {})
    content = {k: v for k, v in cfg.items() if k != "brand"}
    return brand, content


def main(argv=None):
    parser = argparse.ArgumentParser(description="Brand-neutral quote / proposal generator.")
    parser.add_argument("--config", required=True, help="Path to a YAML brand + content config.")
    parser.add_argument("--out", default=None,
                        help="Output .docx path. Default: <config dir>/out/<config stem>.docx")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF rendering.")
    args = parser.parse_args(argv)

    config_path = Path(args.config).resolve()
    if not config_path.exists():
        print(f"config not found: {config_path}")
        return 1

    brand, content = load_config(config_path)
    base_dir = config_path.parent

    if args.out:
        out_path = Path(args.out)
    else:
        out_name = content.get("output", {}).get("docx_name", config_path.stem + ".docx")
        out_path = base_dir / "out" / out_name

    engine = CotizacionEngine(brand, content, base_dir=base_dir)
    print(f"\n=== Building quote from {config_path.name} ===")
    docx_path, pdf_path = engine.build(out_path, make_pdf=not args.no_pdf)
    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
